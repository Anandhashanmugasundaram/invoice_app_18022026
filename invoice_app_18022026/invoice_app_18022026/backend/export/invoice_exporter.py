# """
# Unified Invoice Export System

# Handles PDF and PNG export for any invoice template.
# Replaces the old template-specific export functions.
# """

# import os
# import tempfile
# import logging
# from typing import Tuple, Optional, Literal
# from datetime import datetime
# from playwright.sync_api import sync_playwright
# from templates.invoice_templates import render_invoice_html_template
# from core.invoice_math import calculate_invoice_totals
# from openpyxl import Workbook

# logger = logging.getLogger(__name__)

# # A4 dimensions
# A4_CSS_WIDTH = 794
# A4_CSS_HEIGHT = 1123
# PNG_SCALE_FACTOR = 2480 / A4_CSS_WIDTH  # ~3.125 for ~300dpi


# class InvoiceExporter:
#     """
#     Unified export system for invoice templates.
#     Handles PDF and PNG generation for any template.
#     """
    
#     def __init__(self):
#         self.temp_files = []
    
#     def export(
#         self,
#         invoice_data: dict,
#         template_name: str,
#         format: Literal['pdf', 'png', 'xlsx', 'docx'],
#         output_dir: str = "exports",
#         filename: Optional[str] = None,
#         recalculate: bool = True
#     ) -> Tuple[Optional[str], str]:
#         """
#         Export invoice to PDF or PNG.
        
#         Args:
#             invoice_data: Invoice data dictionary
#             template_name: Template name (e.g., "Siva Sakthi GTA", "Siva Sakthi Freight Bill")
#             format: 'pdf' or 'png'
#             output_dir: Output directory path
#             filename: Optional custom filename (without extension)
#             recalculate: Whether to recalculate totals before export
        
#         Returns:
#             Tuple of (file_path, status_message)
#         """
#         try:
#             # Recalculate totals if requested
#             if recalculate:
#                 invoice_data = calculate_invoice_totals(invoice_data)
            
#             # Create output directory
#             os.makedirs(output_dir, exist_ok=True)
            
#             # Generate filename
#             if not filename:
#                 filename = self._generate_filename(invoice_data)
#             filename = self._safe_filename(filename)
            
#             # Export based on format
#             if format == 'pdf':
#                 return self._export_pdf(invoice_data, template_name, output_dir, filename)
#             elif format == 'png':
#                 return self._export_png(invoice_data, template_name, output_dir, filename)
#             elif format == 'xlsx':
#                  return self._export_excel(invoice_data, output_dir, filename)
#             elif format == 'docx':
#                  return self._export_word(invoice_data, output_dir, filename)
#             else:
#                 return None, f"❌ Unsupported format: {format}"
        
#         except Exception as e:
#             logger.exception(f"Export failed: {e}")
#             return None, f"❌ Export failed: {e}"
    
    

#     def _export_pdf(
#         self,
#         invoice_data: dict,
#         template_name: str,
#         output_dir: str,
#         filename: str
#     ) -> Tuple[Optional[str], str]:
#         """Generate PDF export"""
#         pdf_path = os.path.join(output_dir, f"{filename}.pdf")
#         temp_html = self._create_temp_html(invoice_data, template_name)
        
#         try:
#             with sync_playwright() as p:
#                 browser = p.chromium.launch(headless=True)
#                 page = browser.new_page()
#                 page.goto(f"file://{temp_html}")
#                 page.wait_for_load_state("networkidle")
#                 page.wait_for_timeout(500)
                
#                 # Generate PDF with A4 settings
#                 page.pdf(
#                     path=pdf_path,
#                     format="A4",
#                     print_background=True,
#                     margin={
#                         "top": "0mm",
#                         "right": "0mm",
#                         "bottom": "0mm",
#                         "left": "0mm"
#                     }
#                 )
                
#                 browser.close()
            
#             return pdf_path, f"✓ PDF exported: {os.path.basename(pdf_path)}"
        
#         finally:
#             self._cleanup_temp_file(temp_html)
    
#     def _export_png(
#         self,
#         invoice_data: dict,
#         template_name: str,
#         output_dir: str,
#         filename: str
#     ) -> Tuple[Optional[str], str]:
#         """Generate PNG export"""
#         png_path = os.path.join(output_dir, f"{filename}.png")
#         temp_html = self._create_temp_html(invoice_data, template_name)
        
#         try:
#             with sync_playwright() as p:
#                 browser = p.chromium.launch(headless=True)
                
#                 # Viewport at A4 CSS size; device scale factor yields ~300dpi
#                 context = browser.new_context(
#                     viewport={"width": A4_CSS_WIDTH, "height": A4_CSS_HEIGHT},
#                     device_scale_factor=PNG_SCALE_FACTOR,
#                 )
#                 page = context.new_page()
#                 page.goto(f"file://{temp_html}")
#                 page.wait_for_load_state("networkidle")
#                 page.wait_for_timeout(500)
                
#                 # Ensure A4 dimensions
#                 page.evaluate("""
#                     () => {
#                         const style = document.createElement('style');
#                         style.textContent = `
#                             @page { size: 210mm 297mm; margin: 0; }
#                             body { width: 210mm; height: 297mm; margin: 0; padding: 0; }
#                         `;
#                         document.head.appendChild(style);
#                     }
#                 """)
                
#                 # Take full page screenshot
#                 page.screenshot(
#                     path=png_path,
#                     full_page=True,
#                     type="png"
#                 )
                
#                 browser.close()
            
#             return png_path, f"✓ PNG exported: {os.path.basename(png_path)}"
        
#         finally:
#             self._cleanup_temp_file(temp_html)
    
#     def _create_temp_html(self, invoice_data: dict, template_name: str) -> str:
#         """Create temporary HTML file from invoice data"""
#         html = render_invoice_html_template(invoice_data or {}, template_name)
#         f = tempfile.NamedTemporaryFile(
#             mode="w",
#             suffix=".html",
#             delete=False,
#             encoding="utf-8"
#         )
#         with f:
#             f.write(html)
        
#         self.temp_files.append(f.name)
#         return f.name
    
#     def _cleanup_temp_file(self, filepath: str):
#         """Clean up temporary file"""
#         try:
#             if os.path.exists(filepath):
#                 os.unlink(filepath)
#             if filepath in self.temp_files:
#                 self.temp_files.remove(filepath)
#         except Exception as e:
#             logger.warning(f"Failed to cleanup temp file {filepath}: {e}")
    
#     def _generate_filename(self, invoice_data: dict) -> str:
#         """Generate filename from invoice data"""
#         invoice_no = (invoice_data or {}).get("invoice_number", "").strip()
#         if invoice_no:
#             return f"invoice_{invoice_no}"
#         else:
#             return f"invoice_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    
#     def _safe_filename(self, s: str) -> str:
#         """Sanitize filename"""
#         return "".join(c for c in (s or "") if c.isalnum() or c in ("-", "_")) or "invoice"
    
#     def cleanup_all(self):
#         """Clean up all temporary files"""
#         for filepath in self.temp_files[:]:
#             self._cleanup_temp_file(filepath)


# # Global exporter instance
# _exporter = None

# def get_exporter() -> InvoiceExporter:
#     """Get global InvoiceExporter instance"""
#     global _exporter
#     if _exporter is None:
#         _exporter = InvoiceExporter()
#     return _exporter


# # Convenience functions for direct use
# def export_invoice(
#     invoice_data: dict,
#     template_name: str,
#     format: Literal['pdf', 'png'],
#     output_dir: str = "exports",
#     filename: Optional[str] = None
# ) -> Tuple[Optional[str], str]:
#     """
#     Export invoice to PDF or PNG.
    
#     Args:
#         invoice_data: Invoice data dictionary
#         template_name: Template name (e.g., "Siva Sakthi GTA")
#         format: 'pdf' or 'png'
#         output_dir: Output directory
#         filename: Optional custom filename
    
#     Returns:
#         Tuple of (file_path, status_message)
#     """
#     exporter = get_exporter()
#     return exporter.export(invoice_data, template_name, format, output_dir, filename)

#000000000000000000000000000000000000000000000000000000000

# """
# Unified Invoice Export System

# Handles PDF, PNG, Excel and Word export for any invoice template.
# Replaces the old template-specific export functions.
# """

# import os
# import tempfile
# import logging
# from typing import Tuple, Optional, Literal
# from datetime import datetime
# from playwright.sync_api import sync_playwright
# from templates.invoice_templates import render_invoice_html_template
# from core.invoice_math import calculate_invoice_totals
# from openpyxl import Workbook
# from docx import Document

# logger = logging.getLogger(__name__)

# # A4 dimensions
# A4_CSS_WIDTH = 794
# A4_CSS_HEIGHT = 1123
# PNG_SCALE_FACTOR = 2480 / A4_CSS_WIDTH  # ~300 DPI


# class InvoiceExporter:
#     """
#     Unified export system for invoice templates.
#     Handles PDF, PNG, Excel and Word generation.
#     """

#     def __init__(self):
#         self.temp_files = []

#     def export(
#         self,
#         invoice_data: dict,
#         template_name: str,
#         format: Literal['pdf', 'png', 'xlsx', 'docx'],
#         output_dir: str = "exports",
#         filename: Optional[str] = None,
#         recalculate: bool = True
#     ) -> Tuple[Optional[str], str]:

#         try:
#             if recalculate:
#                 invoice_data = calculate_invoice_totals(invoice_data)

#             os.makedirs(output_dir, exist_ok=True)

#             if not filename:
#                 filename = self._generate_filename(invoice_data)

#             filename = self._safe_filename(filename)

#             if format == 'pdf':
#                 return self._export_pdf(invoice_data, template_name, output_dir, filename)

#             elif format == 'png':
#                 return self._export_png(invoice_data, template_name, output_dir, filename)

#             elif format == 'xlsx':
#                 return self._export_excel(invoice_data, output_dir, filename)

#             elif format == 'docx':
#                 return self._export_word(invoice_data, output_dir, filename)

#             else:
#                 return None, f"❌ Unsupported format: {format}"

#         except Exception as e:
#             logger.exception(f"Export failed: {e}")
#             return None, f"❌ Export failed: {e}"

#     # ---------------- PDF ---------------- #

#     def _export_pdf(self, invoice_data, template_name, output_dir, filename):
#         pdf_path = os.path.join(output_dir, f"{filename}.pdf")
#         temp_html = self._create_temp_html(invoice_data, template_name)

#         try:
#             with sync_playwright() as p:
#                 browser = p.chromium.launch(headless=True)
#                 page = browser.new_page()
#                 page.goto(f"file://{temp_html}")
#                 page.wait_for_load_state("networkidle")

#                 page.pdf(
#                     path=pdf_path,
#                     format="A4",
#                     print_background=True,
#                     margin={"top": "0mm", "right": "0mm",
#                             "bottom": "0mm", "left": "0mm"}
#                 )

#                 browser.close()

#             return pdf_path, f"✓ PDF exported: {os.path.basename(pdf_path)}"

#         finally:
#             self._cleanup_temp_file(temp_html)

#     # ---------------- PNG ---------------- #

#     def _export_png(self, invoice_data, template_name, output_dir, filename):
#         png_path = os.path.join(output_dir, f"{filename}.png")
#         temp_html = self._create_temp_html(invoice_data, template_name)

#         try:
#             with sync_playwright() as p:
#                 browser = p.chromium.launch(headless=True)

#                 context = browser.new_context(
#                     viewport={"width": A4_CSS_WIDTH, "height": A4_CSS_HEIGHT},
#                     device_scale_factor=PNG_SCALE_FACTOR,
#                 )

#                 page = context.new_page()
#                 page.goto(f"file://{temp_html}")
#                 page.wait_for_load_state("networkidle")

#                 page.screenshot(path=png_path, full_page=True)

#                 browser.close()

#             return png_path, f"✓ PNG exported: {os.path.basename(png_path)}"

#         finally:
#             self._cleanup_temp_file(temp_html)

#     # ---------------- EXCEL ---------------- #

#     def _export_excel(self, invoice_data, output_dir, filename):
#         excel_path = os.path.join(output_dir, f"{filename}.xlsx")

#         wb = Workbook()
#         ws = wb.active
#         ws.title = "Invoice"

#         row = 1
#         for key, value in (invoice_data or {}).items():
#             ws.cell(row=row, column=1, value=str(key))
#             ws.cell(row=row, column=2, value=str(value))
#             row += 1

#         wb.save(excel_path)

#         return excel_path, f"✓ Excel exported: {os.path.basename(excel_path)}"

#     # ---------------- WORD ---------------- #

#     def _export_word(self, invoice_data, output_dir, filename):
#         word_path = os.path.join(output_dir, f"{filename}.docx")

#         document = Document()
#         document.add_heading("Invoice", level=1)

#         for key, value in (invoice_data or {}).items():
#             document.add_paragraph(f"{key}: {value}")

#         document.save(word_path)

#         return word_path, f"✓ Word exported: {os.path.basename(word_path)}"

#     # ---------------- HELPERS ---------------- #

#     def _create_temp_html(self, invoice_data, template_name):
#         html = render_invoice_html_template(invoice_data or {}, template_name)
#         f = tempfile.NamedTemporaryFile(
#             mode="w",
#             suffix=".html",
#             delete=False,
#             encoding="utf-8"
#         )
#         with f:
#             f.write(html)

#         self.temp_files.append(f.name)
#         return f.name

#     def _cleanup_temp_file(self, filepath):
#         if os.path.exists(filepath):
#             os.unlink(filepath)

#     def _generate_filename(self, invoice_data):
#         invoice_no = (invoice_data or {}).get("invoice_number", "").strip()
#         if invoice_no:
#             return f"invoice_{invoice_no}"
#         return f"invoice_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

#     def _safe_filename(self, s):
#         return "".join(c for c in (s or "") if c.isalnum() or c in ("-", "_")) or "invoice"


# # ---------------- GLOBAL ACCESS ---------------- #

# _exporter = None


# def get_exporter():
#     global _exporter
#     if _exporter is None:
#         _exporter = InvoiceExporter()
#     return _exporter


# def export_invoice(
#     invoice_data: dict,
#     template_name: str,
#     format: Literal['pdf', 'png', 'xlsx', 'docx'],
#     output_dir: str = "exports",
#     filename: Optional[str] = None
# ):
#     exporter = get_exporter()
#     return exporter.export(invoice_data, template_name, format, output_dir, filename)



# """
# Unified Invoice Export System
# Handles PDF, PNG, Excel and Word export for any invoice template.
# """

# import os
# import tempfile
# import logging
# from typing import Tuple, Optional, Literal
# from datetime import datetime
# from playwright.sync_api import sync_playwright
# from templates.invoice_templates import render_invoice_html_template
# from core.invoice_math import calculate_invoice_totals
# from openpyxl import Workbook
# from openpyxl.styles import Font
# from docx import Document
# from docx.shared import Pt
# from docx.enum.text import WD_ALIGN_PARAGRAPH

# logger = logging.getLogger(__name__)

# A4_CSS_WIDTH = 794
# A4_CSS_HEIGHT = 1123
# PNG_SCALE_FACTOR = 2480 / A4_CSS_WIDTH


# class InvoiceExporter:

#     def export(
#         self,
#         invoice_data: dict,
#         template_name: str,
#         format: Literal['pdf', 'png', 'xlsx', 'docx'],
#         output_dir: str = "exports",
#         filename: Optional[str] = None,
#         recalculate: bool = True
#     ) -> Tuple[Optional[str], str]:

#         try:
#             if recalculate:
#                 invoice_data = calculate_invoice_totals(invoice_data)

#             os.makedirs(output_dir, exist_ok=True)

#             if not filename:
#                 filename = self._generate_filename(invoice_data)

#             filename = self._safe_filename(filename)

#             if format == 'pdf':
#                 return self._export_pdf(invoice_data, template_name, output_dir, filename)

#             elif format == 'png':
#                 return self._export_png(invoice_data, template_name, output_dir, filename)

#             elif format == 'xlsx':
#                 return self._export_excel(invoice_data, output_dir, filename)

#             elif format == 'docx':
#                 return self._export_word(invoice_data, output_dir, filename)

#             else:
#                 return None, f"Unsupported format: {format}"

#         except Exception as e:
#             logger.exception(f"Export failed: {e}")
#             return None, f"Export failed: {e}"

#     # ---------------- PDF ---------------- #

#     def _export_pdf(self, invoice_data, template_name, output_dir, filename):
#         pdf_path = os.path.join(output_dir, f"{filename}.pdf")
#         temp_html = self._create_temp_html(invoice_data, template_name)

#         try:
#             with sync_playwright() as p:
#                 browser = p.chromium.launch(headless=True)
#                 page = browser.new_page()
#                 page.goto(f"file://{temp_html}")
#                 page.wait_for_load_state("networkidle")

#                 page.pdf(
#                     path=pdf_path,
#                     format="A4",
#                     print_background=True,
#                     margin={"top": "0mm", "right": "0mm",
#                             "bottom": "0mm", "left": "0mm"}
#                 )

#                 browser.close()

#             return pdf_path, f"PDF exported: {os.path.basename(pdf_path)}"

#         finally:
#             self._cleanup_temp_file(temp_html)

#     # ---------------- PNG ---------------- #

#     def _export_png(self, invoice_data, template_name, output_dir, filename):
#         png_path = os.path.join(output_dir, f"{filename}.png")
#         temp_html = self._create_temp_html(invoice_data, template_name)

#         try:
#             with sync_playwright() as p:
#                 browser = p.chromium.launch(headless=True)

#                 context = browser.new_context(
#                     viewport={"width": A4_CSS_WIDTH, "height": A4_CSS_HEIGHT},
#                     device_scale_factor=PNG_SCALE_FACTOR,
#                 )

#                 page = context.new_page()
#                 page.goto(f"file://{temp_html}")
#                 page.wait_for_load_state("networkidle")

#                 page.screenshot(path=png_path, full_page=True)

#                 browser.close()

#             return png_path, f"PNG exported: {os.path.basename(png_path)}"

#         finally:
#             self._cleanup_temp_file(temp_html)

#     # ---------------- EXCEL ---------------- #

#     def _export_excel(self, invoice_data, output_dir, filename):
#         excel_path = os.path.join(output_dir, f"{filename}.xlsx")

#         wb = Workbook()
#         ws = wb.active
#         ws.title = "Freight Bill"

#         headers = [
#             "S.No", "Date", "Truck No", "LR No",
#             "D.C Qty (MT)", "GR Qty (MT)",
#             "Rate/MT", "Subtotal"
#         ]

#         ws.append(headers)

#         for col in range(1, len(headers) + 1):
#             ws.cell(row=1, column=col).font = Font(bold=True)

#         runs = invoice_data.get("freight_bill", {}).get("runs", [])

#         for index, run in enumerate(runs, start=1):
#             ws.append([
#                 index,
#                 run.get("date"),
#                 run.get("truck_no"),
#                 run.get("lr_no"),
#                 run.get("dc_qty_mt"),
#                 run.get("gr_qty_mt"),
#                 run.get("rate"),
#                 run.get("line_total"),
#             ])

#         # Totals
#         ws.append([])
#         ws.append(["", "", "", "", "", "", "Grand Total",
#                    invoice_data.get("freight_bill", {}).get("subtotal", 0)])

#         wb.save(excel_path)

#         return excel_path, f"Excel exported: {os.path.basename(excel_path)}"

#     # ---------------- WORD ---------------- #

#     def _export_word(self, invoice_data, output_dir, filename):
#         word_path = os.path.join(output_dir, f"{filename}.docx")

#         document = Document()

#         document.add_heading("Freight Bill", level=1)

#         runs = invoice_data.get("freight_bill", {}).get("runs", [])

#         headers = [
#             "S.No", "Date", "Truck No", "LR No",
#             "D.C Qty (MT)", "GR Qty (MT)",
#             "Rate/MT", "Subtotal"
#         ]

#         table = document.add_table(rows=1, cols=len(headers))
#         table.style = "Table Grid"

#         for i, header in enumerate(headers):
#             table.rows[0].cells[i].text = header

#         for index, run in enumerate(runs, start=1):
#             row = table.add_row().cells
#             row[0].text = str(index)
#             row[1].text = str(run.get("date"))
#             row[2].text = str(run.get("truck_no"))
#             row[3].text = str(run.get("lr_no"))
#             row[4].text = str(run.get("dc_qty_mt"))
#             row[5].text = str(run.get("gr_qty_mt"))
#             row[6].text = str(run.get("rate"))
#             row[7].text = str(run.get("line_total"))

#         document.add_paragraph("")
#         document.add_paragraph(
#             f"Grand Total: {invoice_data.get('freight_bill', {}).get('subtotal', 0)}"
#         ).alignment = WD_ALIGN_PARAGRAPH.RIGHT

#         document.save(word_path)

#         return word_path, f"Word exported: {os.path.basename(word_path)}"

#     # ---------------- HELPERS ---------------- #

#     def _create_temp_html(self, invoice_data, template_name):
#         html = render_invoice_html_template(invoice_data or {}, template_name)
#         f = tempfile.NamedTemporaryFile(
#             mode="w",
#             suffix=".html",
#             delete=False,
#             encoding="utf-8"
#         )
#         with f:
#             f.write(html)
#         return f.name

#     def _cleanup_temp_file(self, filepath):
#         if os.path.exists(filepath):
#             os.unlink(filepath)

#     def _generate_filename(self, invoice_data):
#         invoice_no = (invoice_data or {}).get("invoice_number", "").strip()
#         if invoice_no:
#             return f"invoice_{invoice_no}"
#         return f"invoice_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

#     def _safe_filename(self, s):
#         return "".join(c for c in (s or "") if c.isalnum() or c in ("-", "_")) or "invoice"


# # ---------------- GLOBAL ACCESS ---------------- #

# _exporter = None


# def get_exporter():
#     global _exporter
#     if _exporter is None:
#         _exporter = InvoiceExporter()
#     return _exporter


# def export_invoice(
#     invoice_data: dict,
#     template_name: str,
#     format: Literal['pdf', 'png', 'xlsx', 'docx'],
#     output_dir: str = "exports",
#     filename: Optional[str] = None
# ):
#     exporter = get_exporter()
#     return exporter.export(invoice_data, template_name, format, output_dir, filename)


# """
# Unified Invoice Export System
# Handles PDF, PNG, Excel and Word export for any invoice template.
# FINAL FIXED VERSION (ALL FEATURES WORKING + CORRECT WORD FORMAT)
# """

# import os
# import tempfile
# import logging
# from typing import Tuple, Optional, Literal
# from datetime import datetime

# from playwright.sync_api import sync_playwright
# from templates.invoice_templates import render_invoice_html_template
# from core.invoice_math import calculate_invoice_totals

# from openpyxl import Workbook
# from openpyxl.styles import Font

# from docx import Document
# from docx.enum.text import WD_ALIGN_PARAGRAPH

# logger = logging.getLogger(__name__)

# A4_CSS_WIDTH = 794
# A4_CSS_HEIGHT = 1123
# PNG_SCALE_FACTOR = 2480 / A4_CSS_WIDTH


# class InvoiceExporter:

#     # ==========================================================
#     # MAIN EXPORT
#     # ==========================================================

#     def export(
#         self,
#         invoice_data: dict,
#         template_name: str,
#         format: Literal['pdf', 'png', 'xlsx', 'docx'],
#         output_dir: str = "exports",
#         filename: Optional[str] = None,
#         recalculate: bool = True
#     ) -> Tuple[Optional[str], str]:

#         try:
#             if recalculate:
#                 invoice_data = calculate_invoice_totals(invoice_data)

#             os.makedirs(output_dir, exist_ok=True)

#             if not filename:
#                 filename = self._generate_filename(invoice_data)

#             filename = self._safe_filename(filename)

#             if format == 'pdf':
#                 return self._export_pdf(invoice_data, template_name, output_dir, filename)

#             elif format == 'png':
#                 return self._export_png(invoice_data, template_name, output_dir, filename)

#             elif format == 'xlsx':
#                 return self._export_excel(invoice_data, output_dir, filename)

#             elif format == 'docx':
#                 return self._export_word(invoice_data, output_dir, filename)

#             else:
#                 return None, f"Unsupported format: {format}"

#         except Exception as e:
#             logger.exception(f"Export failed: {e}")
#             return None, f"Export failed: {e}"

#     # ==========================================================
#     # PDF EXPORT
#     # ==========================================================

#     def _export_pdf(self, invoice_data, template_name, output_dir, filename):
#         pdf_path = os.path.join(output_dir, f"{filename}.pdf")
#         temp_html = self._create_temp_html(invoice_data, template_name)

#         try:
#             with sync_playwright() as p:
#                 browser = p.chromium.launch(headless=True)
#                 page = browser.new_page()
#                 page.goto(f"file://{temp_html}")
#                 page.wait_for_load_state("networkidle")

#                 page.pdf(
#                     path=pdf_path,
#                     format="A4",
#                     print_background=True,
#                     margin={"top": "0mm", "right": "0mm",
#                             "bottom": "0mm", "left": "0mm"}
#                 )

#                 browser.close()

#             return pdf_path, f"PDF exported: {os.path.basename(pdf_path)}"

#         finally:
#             self._cleanup_temp_file(temp_html)

#     # ==========================================================
#     # PNG EXPORT (FIXED PROPERLY)
#     # ==========================================================

#     def _export_png(self, invoice_data, template_name, output_dir, filename):
#         png_path = os.path.join(output_dir, f"{filename}.png")
#         temp_html = self._create_temp_html(invoice_data, template_name)

#         try:
#             with sync_playwright() as p:
#                 browser = p.chromium.launch(headless=True)

#                 context = browser.new_context(
#                     viewport={"width": A4_CSS_WIDTH, "height": A4_CSS_HEIGHT},
#                     device_scale_factor=PNG_SCALE_FACTOR,
#                 )

#                 page = context.new_page()
#                 page.goto(f"file://{temp_html}")
#                 page.wait_for_load_state("networkidle")

#                 page.screenshot(path=png_path, full_page=True)

#                 context.close()
#                 browser.close()

#             return png_path, f"PNG exported: {os.path.basename(png_path)}"

#         finally:
#             self._cleanup_temp_file(temp_html)

#             # ==========================================================
#     # WORD EXPORT
#     # ==========================================================
#     def _export_word(self, invoice_data, output_dir, filename):
#         word_path = os.path.join(output_dir, f"{filename}.docx")
#         document = Document()

#         company = invoice_data.get("company_info", {})
#         freight = invoice_data.get("freight_bill", {})
#         runs = freight.get("runs", [])

#         # 1️⃣ HEADER TABLE
#         header_table = document.add_table(rows=1, cols=2)
#         header_table.style = "Table Grid"

#         left_cell = header_table.rows[0].cells[0]
#         right_cell = header_table.rows[0].cells[1]

#         left_cell.text = (
#             f"{company.get('name', 'Siva Sakthi Roadways')}\n"
#             f"GOODS TRANSPORT AGENCY & WEIGH BRIDGE\n"
#             f"{company.get('address', '# D-46, CMDA Truck Terminal Complex, Madhavaram, Chennai - 600 110')}\n"
#             f"Mobile: {company.get('mobile', '094453 84189')}"
#         )

#         right_cell.text = f"PAN NO : {company.get('pan', 'ARMPS3396J')}"
#         right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

#         document.add_paragraph("")

#         # 2️⃣ TO BLOCK
#         to_table = document.add_table(rows=1, cols=1)
#         to_table.style = "Table Grid"

#         to_party = freight.get("to_party", {})

#         to_table.rows[0].cells[0].text = (
#             f"To\n"
#             f"{to_party.get('name', 'ZUARI CEMENT LTD.,')}\n"
#             f"{to_party.get('address', 'CGU - Attipattu,Chennai')}"
#         )

#         document.add_paragraph("")

#         # 3️⃣ BILL INFO
#         bill_table = document.add_table(rows=1, cols=3)
#         bill_table.style = "Table Grid"

#         bill_table.rows[0].cells[0].text = f"Freight Bill {freight.get('series_no', '')}"
#         bill_table.rows[0].cells[1].text = f"Dt {freight.get('bill_date', '')}"
#         bill_table.rows[0].cells[2].text = f"P.O NO : {freight.get('po_no', '')}"

#         document.add_paragraph("")

#         # 4️⃣ TITLE
#         title_para = document.add_paragraph("Transport of Gypsum")
#         title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

#         document.add_paragraph("")

#         # 5️⃣ MAIN TABLE
#         headers = [
#             "S.No", "Date", "Truck No", "LR No",
#             "D.C Qty\nin MT", "GR Qty\nin MT",
#             "Rate/MT", "Subtotal"
#         ]

#         table = document.add_table(rows=1, cols=len(headers))
#         table.style = "Table Grid"

#         for i, header in enumerate(headers):
#             table.rows[0].cells[i].text = header

#         for index, run in enumerate(runs, start=1):
#             row = table.add_row().cells
#             row[0].text = str(index)
#             row[1].text = str(run.get("date", ""))
#             row[2].text = str(run.get("truck_no", ""))
#             row[3].text = str(run.get("lr_no", ""))
#             row[4].text = str(run.get("dc_qty_mt", ""))
#             row[5].text = str(run.get("gr_qty_mt", ""))
#             row[6].text = str(run.get("rate", ""))
#             row[7].text = f"{run.get('line_total', 0):,.2f}"

#         # 6️⃣ TOTAL
#         total_row = table.add_row().cells
#         total_row[6].text = "Grand Total"
#         total_row[7].text = f"{freight.get('subtotal', 0):,.2f}"

#         document.save(word_path)

#         return word_path, f"Word exported: {os.path.basename(word_path)}"
        
#     # ==========================================================
# # EXCEL EXPORT (NORMAL WHITE STYLE)
# # ==========================================================
# def _export_excel(self, invoice_data, output_dir, filename):
#     import os
#     from openpyxl import Workbook
#     from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

#     excel_path = os.path.join(output_dir, f"{filename}.xlsx")

#     wb = Workbook()
#     ws = wb.active
#     ws.title = "Freight Bill"

#     company = invoice_data.get("company_info", {})
#     freight = invoice_data.get("freight_bill", {})
#     runs = freight.get("runs", [])

#     # ===== Styles =====
#     bold = Font(bold=True)
#     black_text = Font(bold=True, color="000000")
#     center = Alignment(horizontal="center", vertical="center")
#     right = Alignment(horizontal="right", vertical="center")
#     left = Alignment(horizontal="left", vertical="center")
#     border = Border(
#         left=Side(style="thin"),
#         right=Side(style="thin"),
#         top=Side(style="thin"),
#         bottom=Side(style="thin")
#     )
#     thin_line = Border(top=Side(style="thin"))
#     white_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")

#     # ================= HEADER =================
#     ws.merge_cells("A1:H1")
#     ws["A1"] = company.get("name", "SIVA SAKTHI ROADWAYS")
#     ws["A1"].font = Font(bold=True, size=16)
#     ws["A1"].alignment = center

#     ws.merge_cells("A2:H2")
#     ws["A2"] = "Authorised Transport Contractor"
#     ws["A2"].alignment = center

#     ws.merge_cells("A3:H3")
#     ws["A3"] = company.get(
#         "address",
#         "No.D-46, CMDA Truck Terminal, Madhavaram, Chennai - 600110"
#     )
#     ws["A3"].alignment = center

#     ws["A4"] = f"PH: {company.get('phone', '')}"
#     ws["A4"].alignment = left

#     ws["A5"] = f"Mobile No: {company.get('mobile', '')}"
#     ws["A5"].alignment = left

#     ws.merge_cells("E4:H4")
#     ws["E4"] = f"PAN NO. {company.get('pan', '')}"
#     ws["E4"].alignment = right
#     ws["E4"].font = bold

#     # ================= TO SECTION =================
#     ws["A7"] = "To"
#     ws["A7"].font = bold

#     to_party = freight.get("to_party", {})
#     ws["A8"] = to_party.get("name", "")
#     ws["A9"] = to_party.get("address", "")

#     # ================= BILL INFO =================
#     ws.merge_cells("A11:H11")
#     ws["A11"] = (
#         f"Freight Bill  {freight.get('series_no', '')} , "
#         f"Dt {freight.get('bill_date', '')} , "
#         f"P.O No : {freight.get('po_no', '')}"
#     )
#     ws["A11"].alignment = left
#     ws["A11"].font = bold

#     ws.merge_cells("A12:H12")
#     ws["A12"] = "Transportation of Goods"
#     ws["A12"].alignment = left

#     # ================= TABLE =================
#     headers = [
#         "S.No", "Date", "Truck No", "LR No",
#         "DC Qty", "GR Qty",
#         "Rate/MT", "Total (Rs)"
#     ]

#     start_row = 14
#     for col, header in enumerate(headers, start=1):
#         cell = ws.cell(row=start_row, column=col)
#         cell.value = header
#         cell.font = bold
#         cell.alignment = center
#         cell.border = border

#     row_cursor = start_row + 1
#     for index, run in enumerate(runs, start=1):
#         values = [
#             index,
#             run.get("date", ""),
#             run.get("truck_no", ""),
#             run.get("lr_no", ""),
#             run.get("dc_qty_mt", ""),
#             run.get("gr_qty_mt", ""),
#             run.get("rate", ""),
#             run.get("line_total", 0),
#         ]
#         for col, value in enumerate(values, start=1):
#             cell = ws.cell(row=row_cursor, column=col)
#             cell.value = value
#             cell.border = border
#             cell.alignment = right if col >= 5 else center
#         row_cursor += 1

#     # ================= GRAND TOTAL =================
#     ws.cell(row=row_cursor, column=6, value="Grand Total").font = bold
#     ws.cell(row=row_cursor, column=6).alignment = right
#     ws.cell(row=row_cursor, column=8, value=freight.get("subtotal", 0)).font = bold
#     ws.cell(row=row_cursor, column=8).alignment = right

#     # ================= AMOUNT IN WORDS =================
#     row_cursor += 2
#     ws.merge_cells(f"A{row_cursor}:H{row_cursor}")
#     ws[f"A{row_cursor}"] = f"Amount in Words : {freight.get('amount_in_words', '')}"
#     ws[f"A{row_cursor}"].alignment = left
#     ws[f"A{row_cursor}"].font = black_text

#     # ================= FOOTER (RIGHT SIGNATURE) =================
#     row_cursor += 2
#     ws.merge_cells(f"E{row_cursor}:H{row_cursor}")
#     ws[f"E{row_cursor}"] = "For Siva Sakthi Roadways"
#     ws[f"E{row_cursor}"].alignment = center
#     ws[f"E{row_cursor}"].font = black_text
#     ws[f"E{row_cursor}"].fill = white_fill
#     ws[f"E{row_cursor}"].border = thin_line

#     row_cursor += 2
#     ws.merge_cells(f"E{row_cursor}:H{row_cursor}")
#     ws[f"E{row_cursor}"] = "Authorised Signatory"
#     ws[f"E{row_cursor}"].alignment = center
#     ws[f"E{row_cursor}"].font = black_text
#     ws[f"E{row_cursor}"].fill = white_fill
#     ws[f"E{row_cursor}"].border = thin_line

#     # ================= COLUMN WIDTH =================
#     widths = [6, 12, 15, 10, 10, 10, 10, 15]
#     for i, width in enumerate(widths, start=1):
#         ws.column_dimensions[chr(64 + i)].width = width

#     # Save workbook
#     wb.save(excel_path)

#     return excel_path, f"Excel exported: {os.path.basename(excel_path)}"

# # ==========================================================
# # HELPERS
# # ==========================================================
# def _create_temp_html(self, invoice_data, template_name):
#     import tempfile
#     html = render_invoice_html_template(invoice_data or {}, template_name)
#     f = tempfile.NamedTemporaryFile(
#         mode="w",
#         suffix=".html",
#         delete=False,
#         encoding="utf-8"
#     )
#     with f:
#         f.write(html)
#     return f.name

# def _cleanup_temp_file(self, filepath):
#     import os
#     if os.path.exists(filepath):
#         os.unlink(filepath)

# def _generate_filename(self, invoice_data):
#     from datetime import datetime
#     invoice_no = (invoice_data or {}).get("invoice_number", "").strip()
#     if invoice_no:
#         return f"invoice_{invoice_no}"
#     return f"invoice_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

# def _safe_filename(self, s):
#     return "".join(c for c in (s or "") if c.isalnum() or c in ("-", "_")) or "invoice"


# """
# Unified Invoice Export System
# Handles PDF, PNG, Excel and Word export for any invoice template.
# FINAL FIXED VERSION (ALL FEATURES WORKING + CORRECT WORD FORMAT)
# """

# import os
# import tempfile
# import logging
# from typing import Tuple, Optional, Literal
# from datetime import datetime

# from playwright.sync_api import sync_playwright
# from templates.invoice_templates import render_invoice_html_template
# from core.invoice_math import calculate_invoice_totals

# from openpyxl import Workbook
# from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

# from docx import Document
# from docx.enum.text import WD_ALIGN_PARAGRAPH

# logger = logging.getLogger(__name__)

# A4_CSS_WIDTH = 794
# A4_CSS_HEIGHT = 1123
# PNG_SCALE_FACTOR = 2480 / A4_CSS_WIDTH


# class InvoiceExporter:

#     # ==========================================================
#     # MAIN EXPORT
#     # ==========================================================
#     def export(
#         self,
#         invoice_data: dict,
#         template_name: str,
#         format: Literal['pdf', 'png', 'xlsx', 'docx'],
#         output_dir: str = "exports",
#         filename: Optional[str] = None,
#         recalculate: bool = True
#     ) -> Tuple[Optional[str], str]:

#         try:
#             if recalculate:
#                 invoice_data = calculate_invoice_totals(invoice_data)

#             os.makedirs(output_dir, exist_ok=True)

#             if not filename:
#                 filename = self._generate_filename(invoice_data)

#             filename = self._safe_filename(filename)

#             if format == 'pdf':
#                 return self._export_pdf(invoice_data, template_name, output_dir, filename)
#             elif format == 'png':
#                 return self._export_png(invoice_data, template_name, output_dir, filename)
#             elif format == 'xlsx':
#                 return self._export_excel(invoice_data, output_dir, filename)
#             elif format == 'docx':
#                 return self._export_word(invoice_data, output_dir, filename)
#             else:
#                 return None, f"Unsupported format: {format}"

#         except Exception as e:
#             logger.exception(f"Export failed: {e}")
#             return None, f"Export failed: {e}"

#     # ==========================================================
#     # PDF EXPORT
#     # ==========================================================
#     def _export_pdf(self, invoice_data, template_name, output_dir, filename):
#         pdf_path = os.path.join(output_dir, f"{filename}.pdf")
#         temp_html = self._create_temp_html(invoice_data, template_name)
#         try:
#             with sync_playwright() as p:
#                 browser = p.chromium.launch(headless=True)
#                 page = browser.new_page()
#                 page.goto(f"file://{temp_html}")
#                 page.wait_for_load_state("networkidle")
#                 page.pdf(
#                     path=pdf_path,
#                     format="A4",
#                     print_background=True,
#                     margin={"top": "0mm", "right": "0mm", "bottom": "0mm", "left": "0mm"}
#                 )
#                 browser.close()
#             return pdf_path, f"PDF exported: {os.path.basename(pdf_path)}"
#         finally:
#             self._cleanup_temp_file(temp_html)

#     # ==========================================================
#     # PNG EXPORT
#     # ==========================================================
#     def _export_png(self, invoice_data, template_name, output_dir, filename):
#         png_path = os.path.join(output_dir, f"{filename}.png")
#         temp_html = self._create_temp_html(invoice_data, template_name)
#         try:
#             with sync_playwright() as p:
#                 browser = p.chromium.launch(headless=True)
#                 context = browser.new_context(
#                     viewport={"width": A4_CSS_WIDTH, "height": A4_CSS_HEIGHT},
#                     device_scale_factor=PNG_SCALE_FACTOR,
#                 )
#                 page = context.new_page()
#                 page.goto(f"file://{temp_html}")
#                 page.wait_for_load_state("networkidle")
#                 page.screenshot(path=png_path, full_page=True)
#                 context.close()
#                 browser.close()
#             return png_path, f"PNG exported: {os.path.basename(png_path)}"
#         finally:
#             self._cleanup_temp_file(temp_html)

#     # ==========================================================
#     # WORD EXPORT
#     # ==========================================================
#     def _export_word(self, invoice_data, output_dir, filename):
#         word_path = os.path.join(output_dir, f"{filename}.docx")
#         document = Document()
#         company = invoice_data.get("company_info", {})
#         freight = invoice_data.get("freight_bill", {})
#         runs = freight.get("runs", [])

#         # HEADER TABLE
#         header_table = document.add_table(rows=1, cols=2)
#         header_table.style = "Table Grid"
#         left_cell = header_table.rows[0].cells[0]
#         right_cell = header_table.rows[0].cells[1]

#         left_cell.text = (
#             f"{company.get('name', 'Siva Sakthi Roadways')}\n"
#             f"GOODS TRANSPORT AGENCY & WEIGH BRIDGE\n"
#             f"{company.get('address', '# D-46, CMDA Truck Terminal Complex, Madhavaram, Chennai - 600 110')}\n"
#             f"Mobile: {company.get('mobile', '094453 84189')}"
#         )
#         right_cell.text = f"PAN NO : {company.get('pan', 'ARMPS3396J')}"
#         right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

#         document.add_paragraph("")

#         # TO BLOCK
#         to_table = document.add_table(rows=1, cols=1)
#         to_table.style = "Table Grid"
#         to_party = freight.get("to_party", {})
#         to_table.rows[0].cells[0].text = (
#             f"To\n"
#             f"{to_party.get('name', 'ZUARI CEMENT LTD.,')}\n"
#             f"{to_party.get('address', 'CGU - Attipattu,Chennai')}"
#         )

#         document.add_paragraph("")

#         # BILL INFO
#         bill_table = document.add_table(rows=1, cols=3)
#         bill_table.style = "Table Grid"
#         bill_table.rows[0].cells[0].text = f"Freight Bill {freight.get('series_no', '')}"
#         bill_table.rows[0].cells[1].text = f"Dt {freight.get('bill_date', '')}"
#         bill_table.rows[0].cells[2].text = f"P.O NO : {freight.get('po_no', '')}"

#         document.add_paragraph("")

#         # TITLE
#         title_para = document.add_paragraph("Transport of Gypsum")
#         title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
#         document.add_paragraph("")

#         # MAIN TABLE
#         headers = [
#             "S.No", "Date", "Truck No", "LR No",
#             "D.C Qty\nin MT", "GR Qty\nin MT",
#             "Rate/MT", "Subtotal"
#         ]
#         table = document.add_table(rows=1, cols=len(headers))
#         table.style = "Table Grid"
#         for i, header in enumerate(headers):
#             table.rows[0].cells[i].text = header

#         for index, run in enumerate(runs, start=1):
#             row = table.add_row().cells
#             row[0].text = str(index)
#             row[1].text = str(run.get("date", ""))
#             row[2].text = str(run.get("truck_no", ""))
#             row[3].text = str(run.get("lr_no", ""))
#             row[4].text = str(run.get("dc_qty_mt", ""))
#             row[5].text = str(run.get("gr_qty_mt", ""))
#             row[6].text = str(run.get("rate", ""))
#             row[7].text = f"{run.get('line_total', 0):,.2f}"

#         # TOTAL
#         total_row = table.add_row().cells
#         total_row[6].text = "Grand Total"
#         total_row[7].text = f"{freight.get('subtotal', 0):,.2f}"

#         document.save(word_path)
#         return word_path, f"Word exported: {os.path.basename(word_path)}"

#     # ==========================================================
#     # EXCEL EXPORT
#     # ==========================================================
#     def _export_excel(self, invoice_data, output_dir, filename):
#         import os
#         from openpyxl import Workbook
#         from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

#         excel_path = os.path.join(output_dir, f"{filename}.xlsx")
#         wb = Workbook()
#         ws = wb.active
#         ws.title = "Freight Bill"

#         company = invoice_data.get("company_info", {})
#         freight = invoice_data.get("freight_bill", {})
#         runs = freight.get("runs", [])

#         # Styles
#         bold = Font(bold=True)
#         black_text = Font(bold=True, color="000000")
#         center = Alignment(horizontal="center", vertical="center")
#         right = Alignment(horizontal="right", vertical="center")
#         left = Alignment(horizontal="left", vertical="center")
#         border = Border(left=Side(style="thin"), right=Side(style="thin"),
#                         top=Side(style="thin"), bottom=Side(style="thin"))
#         thin_line = Border(top=Side(style="thin"))
#         white_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")

#         # HEADER
#         ws.merge_cells("A1:H1")
#         ws["A1"] = company.get("name", "SIVA SAKTHI ROADWAYS")
#         ws["A1"].font = Font(bold=True, size=16)
#         ws["A1"].alignment = center

#         ws.merge_cells("A2:H2")
#         ws["A2"] = "Authorised Transport Contractor"
#         ws["A2"].alignment = center

#         ws.merge_cells("A3:H3")
#         ws["A3"] = company.get("address", "No.D-46, CMDA Truck Terminal, Madhavaram, Chennai - 600110")
#         ws["A3"].alignment = center

#         ws["A4"] = f"PH: {company.get('phone', '')}"
#         ws["A4"].alignment = left

#         ws["A5"] = f"Mobile No: {company.get('mobile', '')}"
#         ws["A5"].alignment = left

#         ws.merge_cells("E4:H4")
#         ws["E4"] = f"PAN NO. {company.get('pan', '')}"
#         ws["E4"].alignment = right
#         ws["E4"].font = bold

#         # TO SECTION
#         ws["A7"] = "To"
#         ws["A7"].font = bold
#         to_party = freight.get("to_party", {})
#         ws["A8"] = to_party.get("name", "")
#         ws["A9"] = to_party.get("address", "")

#         # BILL INFO
#         ws.merge_cells("A11:H11")
#         ws["A11"] = (
#             f"Freight Bill {freight.get('series_no', '')} , "
#             f"Dt {freight.get('bill_date', '')} , "
#             f"P.O No : {freight.get('po_no', '')}"
#         )
#         ws["A11"].alignment = left
#         ws["A11"].font = bold

#         ws.merge_cells("A12:H12")
#         ws["A12"] = "Transportation of Goods"
#         ws["A12"].alignment = left

#         # TABLE
#         headers = ["S.No", "Date", "Truck No", "LR No", "DC Qty", "GR Qty", "Rate/MT", "Total (Rs)"]
#         start_row = 14
#         for col, header in enumerate(headers, start=1):
#             cell = ws.cell(row=start_row, column=col)
#             cell.value = header
#             cell.font = bold
#             cell.alignment = center
#             cell.border = border

#         row_cursor = start_row + 1
#         for index, run in enumerate(runs, start=1):
#             values = [
#                 index,
#                 run.get("date", ""),
#                 run.get("truck_no", ""),
#                 run.get("lr_no", ""),
#                 run.get("dc_qty_mt", ""),
#                 run.get("gr_qty_mt", ""),
#                 run.get("rate", ""),
#                 run.get("line_total", 0),
#             ]
#             for col, value in enumerate(values, start=1):
#                 cell = ws.cell(row=row_cursor, column=col)
#                 cell.value = value
#                 cell.border = border
#                 cell.alignment = right if col >= 5 else center
#             row_cursor += 1

#         # GRAND TOTAL
#         ws.cell(row=row_cursor, column=6, value="Grand Total").font = bold
#         ws.cell(row=row_cursor, column=6).alignment = right
#         ws.cell(row=row_cursor, column=8, value=freight.get("subtotal", 0)).font = bold
#         ws.cell(row=row_cursor, column=8).alignment = right

#         # AMOUNT IN WORDS
#         row_cursor += 2
#         ws.merge_cells(f"A{row_cursor}:H{row_cursor}")
#         ws[f"A{row_cursor}"] = f"Amount in Words : {freight.get('amount_in_words', '')}"
#         ws[f"A{row_cursor}"].alignment = left
#         ws[f"A{row_cursor}"].font = black_text

#         # FOOTER (RIGHT SIGNATURE)
#         row_cursor += 2
#         ws.merge_cells(f"E{row_cursor}:H{row_cursor}")
#         ws[f"E{row_cursor}"] = "For Siva Sakthi Roadways"
#         ws[f"E{row_cursor}"].alignment = center
#         ws[f"E{row_cursor}"].font = black_text
#         ws[f"E{row_cursor}"].fill = white_fill
#         ws[f"E{row_cursor}"].border = thin_line

#         row_cursor += 2
#         ws.merge_cells(f"E{row_cursor}:H{row_cursor}")
#         ws[f"E{row_cursor}"] = "Authorised Signatory"
#         ws[f"E{row_cursor}"].alignment = center
#         ws[f"E{row_cursor}"].font = black_text
#         ws[f"E{row_cursor}"].fill = white_fill
#         ws[f"E{row_cursor}"].border = thin_line

#         # COLUMN WIDTH
#         widths = [6, 12, 15, 10, 10, 10, 10, 15]
#         for i, width in enumerate(widths, start=1):
#             ws.column_dimensions[chr(64 + i)].width = width

#         # Save workbook
#         wb.save(excel_path)
#         return excel_path, f"Excel exported: {os.path.basename(excel_path)}"
    
    

#     # ==========================================================
#     # HELPERS
#     # ==========================================================
#     def _create_temp_html(self, invoice_data, template_name):
#         html = render_invoice_html_template(invoice_data or {}, template_name)
#         f = tempfile.NamedTemporaryFile(mode="w", suffix=".html", delete=False, encoding="utf-8")
#         with f:
#             f.write(html)
#         return f.name

#     def _cleanup_temp_file(self, filepath):
#         if os.path.exists(filepath):
#             os.unlink(filepath)

#     def _generate_filename(self, invoice_data):
#         invoice_no = (invoice_data or {}).get("invoice_number", "").strip()
#         if invoice_no:
#             return f"invoice_{invoice_no}"
#         return f"invoice_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

#     def _safe_filename(self, s):
#         return "".join(c for c in (s or "") if c.isalnum() or c in ("-", "_")) or "invoice"


# # ==========================================================
# # GLOBAL ACCESS
# # ==========================================================
# _exporter = None

# def get_exporter():
#     global _exporter
#     if _exporter is None:
#         _exporter = InvoiceExporter()
#     return _exporter

# def export_invoice(invoice_data: dict, template_name: str, format: Literal['pdf', 'png', 'xlsx', 'docx'],
#                    output_dir: str = "exports", filename: Optional[str] = None):
#     exporter = get_exporter()
#     return exporter.export(invoice_data, template_name, format, output_dir, filename)

# import os
# import tempfile
# import logging
# from typing import Tuple, Optional, Literal
# from datetime import datetime

# from playwright.sync_api import sync_playwright
# from templates.invoice_templates import render_invoice_html_template
# from core.invoice_math import calculate_invoice_totals

# from openpyxl import Workbook
# from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

# from docx import Document
# from docx.enum.text import WD_ALIGN_PARAGRAPH

# logger = logging.getLogger(__name__)

# A4_CSS_WIDTH = 794
# A4_CSS_HEIGHT = 1123
# PNG_SCALE_FACTOR = 2480 / A4_CSS_WIDTH


# class InvoiceExporter:

#     # ==========================================================
#     # MAIN EXPORT
#     # ==========================================================
#     def export(
#         self,
#         invoice_data: dict,
#         template_name: str,
#         format: Literal['pdf', 'png', 'xlsx', 'docx'],
#         output_dir: str = "exports",
#         filename: Optional[str] = None,
#         recalculate: bool = True
#     ) -> Tuple[Optional[str], str]:

#         try:
#             if recalculate:
#                 invoice_data = calculate_invoice_totals(invoice_data)

#             os.makedirs(output_dir, exist_ok=True)

#             if not filename:
#                 filename = self._generate_filename(invoice_data)

#             filename = self._safe_filename(filename)

#             if format == 'pdf':
#                 return self._export_pdf(invoice_data, template_name, output_dir, filename)
#             elif format == 'png':
#                 return self._export_png(invoice_data, template_name, output_dir, filename)
#             elif format == 'xlsx':
#                 # ✅ FIX: Switch based on template
#                 if template_name == "tax_invoice":
#                     return self._export_tax_invoice_excel(invoice_data, output_dir, filename)
#                 else:
#                     return self._export_excel(invoice_data, output_dir, filename)
#             elif format == 'docx':
#                 return self._export_word(invoice_data, output_dir, filename)
#             else:
#                 return None, f"Unsupported format: {format}"

#         except Exception as e:
#             logger.exception(f"Export failed: {e}")
#             return None, f"Export failed: {e}"

#     def _generate_filename(self, invoice_data):
#         """
#         Generate safe filename for invoice export
#         """

#         invoice_number = invoice_data.get("invoice_number", "invoice")
#         invoice_date = invoice_data.get("invoice_date", "")

#     # Clean filename
#         invoice_number = str(invoice_number).replace("/", "-").replace(" ", "_")

#         if invoice_date:
#             safe_date = str(invoice_date).replace("/", "-")
#             return f"{invoice_number}_{safe_date}"
    
#         return invoice_number

#     # ==========================================================
#     # TAX INVOICE EXCEL (NEW FIXED METHOD)
#     # ==========================================================
# from openpyxl.styles import Font, Alignment, Border, Side
# from openpyxl.utils import get_column_letter


# def _export_tax_invoice_excel(self, invoice_data, output_dir, filename):

#     from openpyxl import Workbook
#     import os

#     wb = Workbook()
#     ws = wb.active
#     ws.title = "Tax Invoice"

#     # -------------------------
#     # Styles
#     # -------------------------
#     bold = Font(bold=True, size=11)
#     header_font = Font(bold=True, size=14)
#     center = Alignment(horizontal="center", vertical="center")
#     left = Alignment(horizontal="left", vertical="center")
#     right = Alignment(horizontal="right", vertical="center")
#     thin = Side(style="thin")
#     border = Border(left=thin, right=thin, top=thin, bottom=thin)

#     # Column widths
#     widths = [5, 35, 12, 10, 10, 15, 15]
#     for i, width in enumerate(widths, 1):
#         ws.column_dimensions[get_column_letter(i)].width = width

#     row = 1

#     # ===============================
#     # TITLE
#     # ===============================
#     ws.merge_cells("A1:G1")
#     ws["A1"] = "TAX INVOICE"
#     ws["A1"].font = header_font
#     ws["A1"].alignment = center

#     row += 1

#     ws.merge_cells("A2:G2")
#     ws["G2"] = invoice_data.get("invoice_date", "")
#     ws["G2"].alignment = right

#     row += 2

#     # ===============================
#     # HEADER BOX (Supplier + Invoice Info)
#     # ===============================
#     ws.merge_cells("A4:D7")
#     ws["A4"] = "SIVA SAKTHI TRANSPORT AGENCY"
#     ws["A4"].font = bold
#     ws["A4"].alignment = left

#     ws.merge_cells("E4:G4")
#     ws["E4"] = f"Invoice No : {invoice_data.get('invoice_number', '')}"

#     ws.merge_cells("E5:G5")
#     ws["E5"] = f"Invoice Date : {invoice_data.get('invoice_date', '')}"

#     ws.merge_cells("E6:G6")
#     ws["E6"] = f"Mode of Supply : {invoice_data.get('mode_of_supply', '')}"

#     for r in range(4, 8):
#         for c in range(1, 8):
#             ws.cell(r, c).border = border

#     row = 8

#     # ===============================
#     # BILL TO SECTION
#     # ===============================
#     ws.merge_cells("A8:G8")
#     ws["A8"] = "Bill To"
#     ws["A8"].font = bold

#     row = 9

#     ws.merge_cells("A9:D9")
#     ws["A9"] = f"Name : {invoice_data.get('customer_name', '')}"

#     ws.merge_cells("A10:D10")
#     ws["A10"] = f"Address : {invoice_data.get('customer_address', '')}"

#     ws.merge_cells("A11:D11")
#     ws["A11"] = f"GSTIN : {invoice_data.get('customer_gstin', '')}"

#     ws.merge_cells("E9:G11")

#     for r in range(8, 12):
#         for c in range(1, 8):
#             ws.cell(r, c).border = border

#     row = 12

#     # ===============================
#     # ITEM TABLE HEADER
#     # ===============================
#     headers = [
#         "Sr No",
#         "Name of Product / Service",
#         "HSN/SAC",
#         "Qty",
#         "Rate",
#         "Taxable Value",
#         "Total"
#     ]

#     for col, header in enumerate(headers, 1):
#         ws.cell(row, col).value = header
#         ws.cell(row, col).font = bold
#         ws.cell(row, col).alignment = center
#         ws.cell(row, col).border = border

#     row += 1

#     # ===============================
#     # ITEMS
#     # ===============================
#     items = invoice_data.get("items", [])
#     sr = 1
#     total_taxable = 0

#     for item in items:
#         ws.cell(row, 1).value = sr
#         ws.cell(row, 2).value = item.get("description", "")
#         ws.cell(row, 3).value = item.get("hsn_sac", "")
#         ws.cell(row, 4).value = item.get("qty", "")
#         ws.cell(row, 5).value = item.get("rate", "")
#         ws.cell(row, 6).value = item.get("taxable_value", "")
#         ws.cell(row, 7).value = item.get("total", "")

#         for col in range(1, 8):
#             ws.cell(row, col).border = border

#         total_taxable += float(item.get("taxable_value", 0))
#         sr += 1
#         row += 1

#     # ===============================
#     # TOTAL ROW
#     # ===============================
#     ws.cell(row, 5).value = "Total"
#     ws.cell(row, 5).font = bold
#     ws.cell(row, 6).value = total_taxable
#     ws.cell(row, 6).font = bold

#     for col in range(1, 8):
#         ws.cell(row, col).border = border

#     row += 2

#     # ===============================
#     # TAX SUMMARY BOX
#     # ===============================
#     ws.merge_cells(f"A{row}:D{row}")
#     ws[f"A{row}"] = "Total Invoice Amount in Words:"
#     ws[f"A{row}"].font = bold

#     ws.merge_cells(f"E{row}:G{row}")
#     ws[f"E{row}"] = total_taxable

#     for r in range(row, row + 4):
#         for c in range(1, 8):
#             ws.cell(r, c).border = border

#     row += 5

#     # ===============================
#     # FOOTER
#     # ===============================
#     ws.merge_cells(f"A{row}:G{row}")
#     ws[f"A{row}"] = "Authorised Signatory"
#     ws[f"A{row}"].alignment = right

#     # Save file
#     os.makedirs(output_dir, exist_ok=True)
#     filepath = os.path.join(output_dir, filename)
#     wb.save(filepath)

#     return filepath, "success"

#     # ==========================================================
#     # (YOUR EXISTING FREIGHT EXCEL CODE REMAINS 100% SAME BELOW)
#     # ==========================================================

#         # ==========================================================
#     # FREIGHT EXCEL EXPORT (UNCHANGED)
#     # ==========================================================
#     def _export_excel(self, invoice_data, output_dir, filename):
#         import os
#         from openpyxl import Workbook
#         from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

#         excel_path = os.path.join(output_dir, f"{filename}.xlsx")
#         wb = Workbook()
#         ws = wb.active
#         ws.title = "Freight Bill"

#         company = invoice_data.get("company_info", {})
#         freight = invoice_data.get("freight_bill", {})
#         runs = freight.get("runs", [])

#         # Styles
#         bold = Font(bold=True)
#         black_text = Font(bold=True, color="000000")
#         center = Alignment(horizontal="center", vertical="center")
#         right = Alignment(horizontal="right", vertical="center")
#         left = Alignment(horizontal="left", vertical="center")
#         border = Border(left=Side(style="thin"), right=Side(style="thin"),
#                         top=Side(style="thin"), bottom=Side(style="thin"))
#         thin_line = Border(top=Side(style="thin"))
#         white_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")

#         # HEADER
#         ws.merge_cells("A1:H1")
#         ws["A1"] = company.get("name", "SIVA SAKTHI ROADWAYS")
#         ws["A1"].font = Font(bold=True, size=16)
#         ws["A1"].alignment = center

#         ws.merge_cells("A2:H2")
#         ws["A2"] = "Authorised Transport Contractor"
#         ws["A2"].alignment = center

#         ws.merge_cells("A3:H3")
#         ws["A3"] = company.get("address", "No.D-46, CMDA Truck Terminal, Madhavaram, Chennai - 600110")
#         ws["A3"].alignment = center

#         ws["A4"] = f"PH: {company.get('phone', '')}"
#         ws["A4"].alignment = left

#         ws["A5"] = f"Mobile No: {company.get('mobile', '')}"
#         ws["A5"].alignment = left

#         ws.merge_cells("E4:H4")
#         ws["E4"] = f"PAN NO. {company.get('pan', '')}"
#         ws["E4"].alignment = right
#         ws["E4"].font = bold

#         # TO SECTION
#         ws["A7"] = "To"
#         ws["A7"].font = bold
#         to_party = freight.get("to_party", {})
#         ws["A8"] = to_party.get("name", "")
#         ws["A9"] = to_party.get("address", "")

#         # BILL INFO
#         ws.merge_cells("A11:H11")
#         ws["A11"] = (
#             f"Freight Bill {freight.get('series_no', '')} , "
#             f"Dt {freight.get('bill_date', '')} , "
#             f"P.O No : {freight.get('po_no', '')}"
#         )
#         ws["A11"].alignment = left
#         ws["A11"].font = bold

#         ws.merge_cells("A12:H12")
#         ws["A12"] = "Transportation of Goods"
#         ws["A12"].alignment = left

#         # TABLE
#         headers = ["S.No", "Date", "Truck No", "LR No", "DC Qty", "GR Qty", "Rate/MT", "Total (Rs)"]
#         start_row = 14
#         for col, header in enumerate(headers, start=1):
#             cell = ws.cell(row=start_row, column=col)
#             cell.value = header
#             cell.font = bold
#             cell.alignment = center
#             cell.border = border

#         row_cursor = start_row + 1
#         for index, run in enumerate(runs, start=1):
#             values = [
#                 index,
#                 run.get("date", ""),
#                 run.get("truck_no", ""),
#                 run.get("lr_no", ""),
#                 run.get("dc_qty_mt", ""),
#                 run.get("gr_qty_mt", ""),
#                 run.get("rate", ""),
#                 run.get("line_total", 0),
#             ]
#             for col, value in enumerate(values, start=1):
#                 cell = ws.cell(row=row_cursor, column=col)
#                 cell.value = value
#                 cell.border = border
#                 cell.alignment = right if col >= 5 else center
#             row_cursor += 1

#         # GRAND TOTAL
#         ws.cell(row=row_cursor, column=6, value="Grand Total").font = bold
#         ws.cell(row=row_cursor, column=6).alignment = right
#         ws.cell(row=row_cursor, column=8, value=freight.get("subtotal", 0)).font = bold
#         ws.cell(row=row_cursor, column=8).alignment = right

#         # AMOUNT IN WORDS
#         row_cursor += 2
#         ws.merge_cells(f"A{row_cursor}:H{row_cursor}")
#         ws[f"A{row_cursor}"] = f"Amount in Words : {freight.get('amount_in_words', '')}"
#         ws[f"A{row_cursor}"].alignment = left
#         ws[f"A{row_cursor}"].font = black_text

#         # FOOTER (RIGHT SIGNATURE)
#         row_cursor += 2
#         ws.merge_cells(f"E{row_cursor}:H{row_cursor}")
#         ws[f"E{row_cursor}"] = "For Siva Sakthi Roadways"
#         ws[f"E{row_cursor}"].alignment = center
#         ws[f"E{row_cursor}"].font = black_text
#         ws[f"E{row_cursor}"].fill = white_fill
#         ws[f"E{row_cursor}"].border = thin_line

#         row_cursor += 2
#         ws.merge_cells(f"E{row_cursor}:H{row_cursor}")
#         ws[f"E{row_cursor}"] = "Authorised Signatory"
#         ws[f"E{row_cursor}"].alignment = center
#         ws[f"E{row_cursor}"].font = black_text
#         ws[f"E{row_cursor}"].fill = white_fill
#         ws[f"E{row_cursor}"].border = thin_line

#         # COLUMN WIDTH
#         widths = [6, 12, 15, 10, 10, 10, 10, 15]
#         for i, width in enumerate(widths, start=1):
#             ws.column_dimensions[chr(64 + i)].width = width

#         wb.save(excel_path)
#         return excel_path, f"Excel exported: {os.path.basename(excel_path)}"
    
#         # ==========================================================
#     # HELPERS
#     # ==========================================================
#     def _create_temp_html(self, invoice_data, template_name):
#         html = render_invoice_html_template(invoice_data or {}, template_name)
#         f = tempfile.NamedTemporaryFile(mode="w", suffix=".html", delete=False, encoding="utf-8")
#         with f:
#             f.write(html)
#         return f.name

#     def _cleanup_temp_file(self, filepath):
#         if os.path.exists(filepath):
#             os.unlink(filepath)

#     def _generate_filename(self, invoice_data):
#         invoice_no = (invoice_data or {}).get("invoice_number", "").strip()
#         if invoice_no:
#             return f"invoice_{invoice_no}"
#         return f"invoice_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

#     def _safe_filename(self, s):
#         return "".join(c for c in (s or "") if c.isalnum() or c in ("-", "_")) or "invoice"

#     # ==========================================================
# # GLOBAL ACCESS
# # ==========================================================
# _exporter = None

# def get_exporter():
#     global _exporter
#     if _exporter is None:
#         _exporter = InvoiceExporter()
#     return _exporter


# def export_invoice(
#     invoice_data: dict,
#     template_name: str,
#     format: Literal['pdf', 'png', 'xlsx', 'docx'],
#     output_dir: str = "exports",
#     filename: Optional[str] = None
# ):
#     exporter = get_exporter()
#     return exporter.export(invoice_data, template_name, format, output_dir, filename)


# import os
# import tempfile
# import logging
# from typing import Tuple, Optional, Literal
# from datetime import datetime

# from playwright.sync_api import sync_playwright
# from templates.invoice_templates import render_invoice_html_template
# from core.invoice_math import calculate_invoice_totals

# from openpyxl import Workbook
# from openpyxl.styles import Font, Alignment, Border, Side
# from openpyxl.utils import get_column_letter

# from docx import Document
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.shared import Pt

# logger = logging.getLogger(__name__)

# A4_CSS_WIDTH = 794
# A4_CSS_HEIGHT = 1123
# PNG_SCALE_FACTOR = 2480 / A4_CSS_WIDTH


# class InvoiceExporter:

#     # ==========================================================
#     # MAIN EXPORT
#     # ==========================================================
#     def export(
#         self,
#         invoice_data: dict,
#         template_name: str,
#         format: Literal['pdf', 'png', 'xlsx', 'docx'],
#         output_dir: str = "exports",
#         filename: Optional[str] = None,
#         recalculate: bool = True
#     ) -> Tuple[Optional[str], str]:

#         try:
#             if recalculate:
#                 invoice_data = calculate_invoice_totals(invoice_data)

#             os.makedirs(output_dir, exist_ok=True)

#             if not filename:
#                 filename = self._generate_filename(invoice_data)

#             filename = self._safe_filename(filename)

#             if format == "pdf":
#                 return self._export_pdf(invoice_data, template_name, output_dir, filename)

#             elif format == "png":
#                 return self._export_png(invoice_data, template_name, output_dir, filename)

#             elif format == "xlsx":
#                 if template_name == "tax_invoice":
#                     return self._export_tax_invoice_excel(invoice_data, output_dir, filename)
#                 return self._export_excel(invoice_data, output_dir, filename)

#             elif format == "docx":
#                 return self._export_word(invoice_data, template_name, output_dir, filename)

#             return None, f"Unsupported format: {format}"

#         except Exception as e:
#             logger.exception("Export failed")
#             return None, f"Export failed: {e}"

#     # ==========================================================
#     # PDF EXPORT
#     # ==========================================================
#     def _export_pdf(self, invoice_data, template_name, output_dir, filename):

#         html_path = self._create_temp_html(invoice_data, template_name)
#         pdf_path = os.path.join(output_dir, f"{filename}.pdf")

#         try:
#             with sync_playwright() as p:
#                 browser = p.chromium.launch()
#                 page = browser.new_page()
#                 page.goto(f"file://{html_path}")
#                 page.pdf(
#                     path=pdf_path,
#                     format="A4",
#                     print_background=True
#                 )
#                 browser.close()
#         finally:
#             self._cleanup_temp_file(html_path)

#         return pdf_path, "PDF exported successfully"

#     # ==========================================================
#     # PNG EXPORT
#     # ==========================================================
# def _export_tax_invoice_excel(self, invoice_data, output_dir, filename):

#     wb = Workbook()
#     ws = wb.active
#     ws.title = "Tax Invoice"

#     bold = Font(bold=True)
#     header_font = Font(bold=True, size=14)
#     center = Alignment(horizontal="center", vertical="center")
#     left = Alignment(horizontal="left", vertical="center")
#     right = Alignment(horizontal="right", vertical="center")

#     thin = Side(style="thin")
#     border = Border(left=thin, right=thin, top=thin, bottom=thin)

#     # Column widths
#     widths = [6, 35, 12, 10, 12, 15, 15]
#     for i, width in enumerate(widths, 1):
#         ws.column_dimensions[get_column_letter(i)].width = width

#     # ===============================
#     # TITLE
#     # ===============================
#     ws.merge_cells("A1:G1")
#     ws["A1"] = "TAX INVOICE"
#     ws["A1"].font = header_font
#     ws["A1"].alignment = center

#     # ===============================
#     # COMPANY + INVOICE INFO
#     # ===============================
#     ws.merge_cells("A3:D6")
#     ws["A3"] = "SIVA SAKTHI TRANSPORT AGENCY"
#     ws["A3"].font = bold
#     ws["A3"].alignment = left

#     ws.merge_cells("E3:G3")
#     ws["E3"] = f"Invoice No : {invoice_data.get('invoice_number','')}"

#     ws.merge_cells("E4:G4")
#     ws["E4"] = f"Invoice Date : {invoice_data.get('invoice_date','')}"

#     ws.merge_cells("E5:G5")
#     ws["E5"] = f"Mode of Supply : {invoice_data.get('mode_of_supply','')}"

#     for r in range(3, 7):
#         for c in range(1, 8):
#             ws.cell(r, c).border = border

#     # ===============================
#     # BILL TO
#     # ===============================
#     ws.merge_cells("A7:G7")
#     ws["A7"] = "Bill To"
#     ws["A7"].font = bold

#     ws.merge_cells("A8:D8")
#     ws["A8"] = f"Name : {invoice_data.get('customer_name','')}"

#     ws.merge_cells("A9:D9")
#     ws["A9"] = f"Address : {invoice_data.get('customer_address','')}"

#     ws.merge_cells("A10:D10")
#     ws["A10"] = f"GSTIN : {invoice_data.get('customer_gstin','')}"

#     for r in range(7, 11):
#         for c in range(1, 8):
#             ws.cell(r, c).border = border

#     # ===============================
#     # TABLE HEADER
#     # ===============================
#     headers = ["Sr No", "Description", "HSN/SAC", "Qty", "Rate", "Taxable", "Total"]

#     row = 12
#     for col, header in enumerate(headers, 1):
#         ws.cell(row, col).value = header
#         ws.cell(row, col).font = bold
#         ws.cell(row, col).alignment = center
#         ws.cell(row, col).border = border

#     row += 1
#     total_taxable = 0
#     items = invoice_data.get("items", [])

#     for i, item in enumerate(items, 1):
#         ws.cell(row, 1).value = i
#         ws.cell(row, 2).value = item.get("description","")
#         ws.cell(row, 3).value = item.get("hsn_sac","")
#         ws.cell(row, 4).value = item.get("qty","")
#         ws.cell(row, 5).value = item.get("rate","")
#         ws.cell(row, 6).value = item.get("taxable_value","")
#         ws.cell(row, 7).value = item.get("total","")

#         for c in range(1,8):
#             ws.cell(row, c).border = border

#         total_taxable += float(item.get("taxable_value",0))
#         row += 1

#     # ===============================
#     # TOTAL
#     # ===============================
#     ws.cell(row,5).value = "Total"
#     ws.cell(row,5).font = bold
#     ws.cell(row,5).alignment = right

#     ws.cell(row,6).value = total_taxable
#     ws.cell(row,6).font = bold

#     for c in range(1,8):
#         ws.cell(row,c).border = border

#     # ===============================
#     # SIGNATURE
#     # ===============================
#     row += 3
#     ws.merge_cells(f"E{row}:G{row}")
#     ws[f"E{row}"] = "Authorised Signatory"
#     ws[f"E{row}"].alignment = center
#     ws[f"E{row}"].font = bold

#     path = os.path.join(output_dir, f"{filename}.xlsx")
#     wb.save(path)

#     return path, "Tax invoice Excel exported"

#     # ==========================================================
#     # SIMPLE FREIGHT EXCEL
#     # ==========================================================def _export_excel(self, invoice_data, output_dir, filename):

#     wb = Workbook()
#     ws = wb.active
#     ws.title = "Freight Bill"

#     bold = Font(bold=True)
#     center = Alignment(horizontal="center", vertical="center")
#     right = Alignment(horizontal="right", vertical="center")
#     border = Border(
#         left=Side(style="thin"),
#         right=Side(style="thin"),
#         top=Side(style="thin"),
#         bottom=Side(style="thin")
#     )

#     # Column Widths
#     widths = [6, 12, 15, 12, 10, 10, 12, 15]
#     for i, width in enumerate(widths, 1):
#         ws.column_dimensions[get_column_letter(i)].width = width

#     # ===============================
#     # HEADER
#     # ===============================
#     ws.merge_cells("A1:H1")
#     ws["A1"] = "SIVA SAKTHI ROADWAYS"
#     ws["A1"].font = Font(bold=True, size=16)
#     ws["A1"].alignment = center

#     ws.merge_cells("A2:H2")
#     ws["A2"] = "Freight Bill"
#     ws["A2"].alignment = center

#     # ===============================
#     # TABLE HEADER
#     # ===============================
#     headers = ["S.No", "Date", "Truck No", "LR No", "DC Qty", "GR Qty", "Rate", "Total"]

#     row = 4
#     for col, header in enumerate(headers,1):
#         ws.cell(row,col).value = header
#         ws.cell(row,col).font = bold
#         ws.cell(row,col).alignment = center
#         ws.cell(row,col).border = border

#     # ===============================
#     # DATA
#     # ===============================
#     runs = invoice_data.get("freight_bill", {}).get("runs", [])
#     row += 1

#     for i, run in enumerate(runs,1):
#         values = [
#             i,
#             run.get("date",""),
#             run.get("truck_no",""),
#             run.get("lr_no",""),
#             run.get("dc_qty_mt",""),
#             run.get("gr_qty_mt",""),
#             run.get("rate",""),
#             run.get("line_total","")
#         ]

#         for col,val in enumerate(values,1):
#             ws.cell(row,col).value = val
#             ws.cell(row,col).border = border
#             ws.cell(row,col).alignment = right if col >=5 else center

#         row += 1

#     # ===============================
#     # GRAND TOTAL
#     # ===============================
#     ws.cell(row,7).value = "Grand Total"
#     ws.cell(row,7).font = bold
#     ws.cell(row,7).alignment = right

#     ws.cell(row,8).value = invoice_data.get("freight_bill",{}).get("subtotal",0)
#     ws.cell(row,8).font = bold

#     path = os.path.join(output_dir, f"{filename}.xlsx")
#     wb.save(path)

#     return path, "Freight Excel exported"

#     # ==========================================================
#     # HELPERS
#     # ==========================================================
#     def _create_temp_html(self, invoice_data, template_name):
#         html = render_invoice_html_template(invoice_data or {}, template_name)
#         f = tempfile.NamedTemporaryFile(mode="w", suffix=".html", delete=False, encoding="utf-8")
#         with f:
#             f.write(html)
#         return f.name

#     def _cleanup_temp_file(self, filepath):
#         if os.path.exists(filepath):
#             os.unlink(filepath)

#     def _generate_filename(self, invoice_data):
#         invoice_no = (invoice_data or {}).get("invoice_number", "").strip()
#         if invoice_no:
#             return f"invoice_{invoice_no}"
#         return f"invoice_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

#     def _safe_filename(self, s):
#         return "".join(c for c in (s or "") if c.isalnum() or c in ("-", "_")) or "invoice"


# # ==========================================================
# # GLOBAL ACCESS
# # ==========================================================
# _exporter = None

# def get_exporter():
#     global _exporter
#     if _exporter is None:
#         _exporter = InvoiceExporter()
#     return _exporter


# def export_invoice(
#     invoice_data: dict,
#     template_name: str,
#     format: Literal['pdf', 'png', 'xlsx', 'docx'],
#     output_dir: str = "exports",
#     filename: Optional[str] = None
# ):
#     exporter = get_exporter()
#     return exporter.export(invoice_data, template_name, format, output_dir, filename)

import os
import tempfile
import logging
from typing import Tuple, Optional, Literal
from datetime import datetime

from playwright.sync_api import sync_playwright
from templates.invoice_templates import render_invoice_html_template
from core.invoice_math import calculate_invoice_totals

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

logger = logging.getLogger(__name__)

A4_CSS_WIDTH = 794
A4_CSS_HEIGHT = 1123
PNG_SCALE_FACTOR = 2480 / A4_CSS_WIDTH


class InvoiceExporter:

    # ==========================================================
    # MAIN EXPORT
    # ==========================================================
    def export(
        self,
        invoice_data: dict,
        template_name: str,
        format: Literal['pdf', 'png', 'xlsx', 'docx'],
        output_dir: str = "exports",
        filename: Optional[str] = None,
        recalculate: bool = True
    ) -> Tuple[Optional[str], str]:

        try:
            if recalculate:
                invoice_data = calculate_invoice_totals(invoice_data)

            os.makedirs(output_dir, exist_ok=True)

            if not filename:
                filename = self._generate_filename(invoice_data)

            filename = self._safe_filename(filename)

            if format == "pdf":
                return self._export_pdf(invoice_data, template_name, output_dir, filename)

            elif format == "png":
                return self._export_png(invoice_data, template_name, output_dir, filename)

            elif format == "xlsx":
                if template_name == "tax_invoice":
                    return self._export_tax_invoice_excel(invoice_data, output_dir, filename)
                return self._export_excel(invoice_data, output_dir, filename)

            elif format == "docx":
                return self._export_word(invoice_data, template_name, output_dir, filename)

            return None, f"Unsupported format: {format}"

        except Exception as e:
            logger.exception("Export failed")
            return None, f"Export failed: {e}"

    # ==========================================================
    # PDF EXPORT
    # ==========================================================
    def _export_pdf(self, invoice_data, template_name, output_dir, filename):

        html_path = self._create_temp_html(invoice_data, template_name)
        pdf_path = os.path.join(output_dir, f"{filename}.pdf")

        try:
            with sync_playwright() as p:
                browser = p.chromium.launch()
                page = browser.new_page()
                page.goto(f"file://{html_path}")
                page.pdf(
                    path=pdf_path,
                    format="A4",
                    print_background=True
                )
                browser.close()
        finally:
            self._cleanup_temp_file(html_path)

        return pdf_path, "PDF exported successfully"

    # ==========================================================
    # PNG EXPORT
    # ==========================================================
    def _export_png(self, invoice_data, template_name, output_dir, filename):

        html_path = self._create_temp_html(invoice_data, template_name)
        png_path = os.path.join(output_dir, f"{filename}.png")

        try:
            with sync_playwright() as p:
                browser = p.chromium.launch()
                page = browser.new_page()
                page.goto(f"file://{html_path}")
                page.set_viewport_size({"width": A4_CSS_WIDTH, "height": A4_CSS_HEIGHT})
                page.screenshot(
                    path=png_path,
                    full_page=True
                )
                browser.close()
        finally:
            self._cleanup_temp_file(html_path)

        return png_path, "PNG exported successfully"

    # ==========================================================
    # WORD EXPORT
    # ==========================================================
    def _export_word(self, invoice_data, template_name, output_dir, filename):
        doc = Document()
        
        # Add title
        title = doc.add_heading('Invoice', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add invoice details
        doc.add_paragraph(f"Invoice Number: {invoice_data.get('invoice_number', '')}")
        doc.add_paragraph(f"Invoice Date: {invoice_data.get('invoice_date', '')}")
        doc.add_paragraph(f"Customer: {invoice_data.get('customer_name', '')}")
        
        # Add items table (basic implementation)
        items = invoice_data.get('items', [])
        if items:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Description'
            hdr_cells[1].text = 'Quantity'
            hdr_cells[2].text = 'Rate'
            hdr_cells[3].text = 'Total'
            
            for item in items:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item.get('description', ''))
                row_cells[1].text = str(item.get('qty', ''))
                row_cells[2].text = str(item.get('rate', ''))
                row_cells[3].text = str(item.get('total', ''))
        
        docx_path = os.path.join(output_dir, f"{filename}.docx")
        doc.save(docx_path)
        
        return docx_path, "Word document exported successfully"

    # ==========================================================
    # TAX INVOICE EXCEL
    # ==========================================================
    def _export_tax_invoice_excel(self, invoice_data, output_dir, filename):

        wb = Workbook()
        ws = wb.active
        ws.title = "Tax Invoice"

        bold = Font(bold=True)
        header_font = Font(bold=True, size=14)
        center = Alignment(horizontal="center", vertical="center")
        left = Alignment(horizontal="left", vertical="center")
        right = Alignment(horizontal="right", vertical="center")

        thin = Side(style="thin")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        # Column widths
        widths = [6, 35, 12, 10, 12, 15, 15]
        for i, width in enumerate(widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = width

        # ===============================
        # TITLE
        # ===============================
        ws.merge_cells("A1:G1")
        ws["A1"] = "TAX INVOICE"
        ws["A1"].font = header_font
        ws["A1"].alignment = center

        # ===============================
        # COMPANY + INVOICE INFO
        # ===============================
        ws.merge_cells("A3:D6")
        ws["A3"] = "SIVA SAKTHI TRANSPORT AGENCY"
        ws["A3"].font = bold
        ws["A3"].alignment = left

        ws.merge_cells("E3:G3")
        ws["E3"] = f"Invoice No : {invoice_data.get('invoice_number','')}"

        ws.merge_cells("E4:G4")
        ws["E4"] = f"Invoice Date : {invoice_data.get('invoice_date','')}"

        ws.merge_cells("E5:G5")
        ws["E5"] = f"Mode of Supply : {invoice_data.get('mode_of_supply','')}"

        for r in range(3, 7):
            for c in range(1, 8):
                ws.cell(r, c).border = border

        # ===============================
        # BILL TO
        # ===============================
        ws.merge_cells("A7:G7")
        ws["A7"] = "Bill To"
        ws["A7"].font = bold

        ws.merge_cells("A8:D8")
        ws["A8"] = f"Name : {invoice_data.get('customer_name','')}"

        ws.merge_cells("A9:D9")
        ws["A9"] = f"Address : {invoice_data.get('customer_address','')}"

        ws.merge_cells("A10:D10")
        ws["A10"] = f"GSTIN : {invoice_data.get('customer_gstin','')}"

        for r in range(7, 11):
            for c in range(1, 8):
                ws.cell(r, c).border = border

        # ===============================
        # TABLE HEADER
        # ===============================
        headers = ["Sr No", "Description", "HSN/SAC", "Qty", "Rate", "Taxable", "Total"]

        row = 12
        for col, header in enumerate(headers, 1):
            ws.cell(row, col).value = header
            ws.cell(row, col).font = bold
            ws.cell(row, col).alignment = center
            ws.cell(row, col).border = border

        row += 1
        total_taxable = 0
        items = invoice_data.get("items", [])

        for i, item in enumerate(items, 1):
            ws.cell(row, 1).value = i
            ws.cell(row, 2).value = item.get("description","")
            ws.cell(row, 3).value = item.get("hsn_sac","")
            ws.cell(row, 4).value = item.get("qty","")
            ws.cell(row, 5).value = item.get("rate","")
            ws.cell(row, 6).value = item.get("taxable_value","")
            ws.cell(row, 7).value = item.get("total","")

            for c in range(1,8):
                ws.cell(row, c).border = border

            total_taxable += float(item.get("taxable_value",0))
            row += 1

        # ===============================
        # TOTAL
        # ===============================
        ws.cell(row,5).value = "Total"
        ws.cell(row,5).font = bold
        ws.cell(row,5).alignment = right

        ws.cell(row,6).value = total_taxable
        ws.cell(row,6).font = bold

        for c in range(1,8):
            ws.cell(row,c).border = border

        # ===============================
        # SIGNATURE
        # ===============================
        row += 3
        ws.merge_cells(f"E{row}:G{row}")
        ws[f"E{row}"] = "Authorised Signatory"
        ws[f"E{row}"].alignment = center
        ws[f"E{row}"].font = bold

        path = os.path.join(output_dir, f"{filename}.xlsx")
        wb.save(path)

        return path, "Tax invoice Excel exported"
      
          # ==========================================================
    # SIMPLE FREIGHT EXCEL (UPDATED WITH COMPANY DETAILS)
    # ==========================================================
    def _export_excel(self, invoice_data, output_dir, filename):

        wb = Workbook()
        ws = wb.active
        ws.title = "Freight Bill"

        bold = Font(bold=True)
        header_font = Font(bold=True, size=16)

        center = Alignment(horizontal="center", vertical="center")
        left = Alignment(horizontal="left", vertical="center")
        right = Alignment(horizontal="right", vertical="center")

        border = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin")
        )

        widths = [6, 12, 15, 12, 10, 10, 12, 15]
        for i, width in enumerate(widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = width

        parent = invoice_data.get("company_info", {})
        client = invoice_data.get("client_info", {})

        # Company Header
        ws.merge_cells("A1:H1")
        ws["A1"] = parent.get("name", "SIVA SAKTHI ROADWAYS")
        ws["A1"].font = header_font
        ws["A1"].alignment = center

        ws.merge_cells("A2:H2")
        ws["A2"] = parent.get("address", "")
        ws["A2"].alignment = center

        ws.merge_cells("A3:H3")
        ws["A3"] = f"GSTIN: {parent.get('gstin','')} | Phone: {parent.get('phone','')}"
        ws["A3"].alignment = center

        # Title
        ws.merge_cells("A5:H5")
        ws["A5"] = "FREIGHT BILL"
        ws["A5"].font = Font(bold=True, size=14)
        ws["A5"].alignment = center

        # Client Details
        ws.merge_cells("A7:D7")
        ws["A7"] = "Bill To:"
        ws["A7"].font = bold

        ws.merge_cells("A8:D8")
        ws["A8"] = client.get("name", "")
        ws["A8"].font = bold

        ws.merge_cells("A9:D9")
        ws["A9"] = client.get("address", "")

        ws.merge_cells("A10:D10")
        ws["A10"] = f"GSTIN: {client.get('gstin','')}"

        ws.merge_cells("E7:H7")
        ws["E7"] = f"Invoice No : {invoice_data.get('invoice_number','')}"

        ws.merge_cells("E8:H8")
        ws["E8"] = f"Invoice Date : {invoice_data.get('invoice_date','')}"

        for r in range(7, 11):
            for c in range(1, 9):
                ws.cell(r, c).border = border

        # Table Header
        headers = ["S.No", "Date", "Truck No", "LR No",
                   "DC Qty", "GR Qty", "Rate", "Total"]

        row = 12
        for col, header in enumerate(headers, 1):
            ws.cell(row, col).value = header
            ws.cell(row, col).font = bold
            ws.cell(row, col).alignment = center
            ws.cell(row, col).border = border

        # Data
        runs = invoice_data.get("freight_bill", {}).get("runs", [])
        row += 1

        for i, run in enumerate(runs, 1):
            values = [
                i,
                run.get("date", ""),
                run.get("truck_no", ""),
                run.get("lr_no", ""),
                run.get("dc_qty_mt", ""),
                run.get("gr_qty_mt", ""),
                run.get("rate", ""),
                run.get("line_total", "")
            ]

            for col, val in enumerate(values, 1):
                ws.cell(row, col).value = val
                ws.cell(row, col).border = border
                ws.cell(row, col).alignment = right if col >= 5 else center

            row += 1

        # Grand Total
        ws.cell(row, 7).value = "Grand Total"
        ws.cell(row, 7).font = bold
        ws.cell(row, 7).alignment = right
        ws.cell(row, 7).border = border

        ws.cell(row, 8).value = invoice_data.get("freight_bill", {}).get("subtotal", 0)
        ws.cell(row, 8).font = bold
        ws.cell(row, 8).border = border

        path = os.path.join(output_dir, f"{filename}.xlsx")
        wb.save(path)

        return path, "Freight Excel exported successfully"

        wb = Workbook()
        ws = wb.active
        ws.title = "Freight Bill"

        bold = Font(bold=True)
        center = Alignment(horizontal="center", vertical="center")
        right = Alignment(horizontal="right", vertical="center")
        border = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin")
        )

        # Column Widths
        widths = [6, 12, 15, 12, 10, 10, 12, 15]
        for i, width in enumerate(widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = width

        # ===============================
        # HEADER
        # ===============================
        ws.merge_cells("A1:H1")
        ws["A1"] = "SIVA SAKTHI ROADWAYS"
        ws["A1"].font = Font(bold=True, size=16)
        ws["A1"].alignment = center

        ws.merge_cells("A2:H2")
        ws["A2"] = "Freight Bill"
        ws["A2"].alignment = center

        # ===============================
        # TABLE HEADER
        # ===============================
        headers = ["S.No", "Date", "Truck No", "LR No", "DC Qty", "GR Qty", "Rate", "Total"]

        row = 4
        for col, header in enumerate(headers,1):
            ws.cell(row,col).value = header
            ws.cell(row,col).font = bold
            ws.cell(row,col).alignment = center
            ws.cell(row,col).border = border

        # ===============================
        # DATA
        # ===============================
        runs = invoice_data.get("freight_bill", {}).get("runs", [])
        row += 1

        for i, run in enumerate(runs,1):
            values = [
                i,
                run.get("date",""),
                run.get("truck_no",""),
                run.get("lr_no",""),
                run.get("dc_qty_mt",""),
                run.get("gr_qty_mt",""),
                run.get("rate",""),
                run.get("line_total","")
            ]

            for col,val in enumerate(values,1):
                ws.cell(row,col).value = val
                ws.cell(row,col).border = border
                ws.cell(row,col).alignment = right if col >=5 else center

            row += 1

        # ===============================
        # GRAND TOTAL
        # ===============================
        ws.cell(row,7).value = "Grand Total"
        ws.cell(row,7).font = bold
        ws.cell(row,7).alignment = right

        ws.cell(row,8).value = invoice_data.get("freight_bill",{}).get("subtotal",0)
        ws.cell(row,8).font = bold

        path = os.path.join(output_dir, f"{filename}.xlsx")
        wb.save(path)

        return path, "Freight Excel exported"

    # ==========================================================
    # HELPERS
    # ==========================================================
    def _create_temp_html(self, invoice_data, template_name):
        html = render_invoice_html_template(invoice_data or {}, template_name)
        f = tempfile.NamedTemporaryFile(mode="w", suffix=".html", delete=False, encoding="utf-8")
        with f:
            f.write(html)
        return f.name

    def _cleanup_temp_file(self, filepath):
        if os.path.exists(filepath):
            os.unlink(filepath)

    def _generate_filename(self, invoice_data):
        invoice_no = (invoice_data or {}).get("invoice_number", "").strip()
        if invoice_no:
            return f"invoice_{invoice_no}"
        return f"invoice_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

    def _safe_filename(self, s):
        return "".join(c for c in (s or "") if c.isalnum() or c in ("-", "_")) or "invoice"


# ==========================================================
# GLOBAL ACCESS
# ==========================================================
_exporter = None

def get_exporter():
    global _exporter
    if _exporter is None:
        _exporter = InvoiceExporter()
    return _exporter


def export_invoice(
    invoice_data: dict,
    template_name: str,
    format: Literal['pdf', 'png', 'xlsx', 'docx'],
    output_dir: str = "exports",
    filename: Optional[str] = None
):
    exporter = get_exporter()
    return exporter.export(invoice_data, template_name, format, output_dir, filename)
